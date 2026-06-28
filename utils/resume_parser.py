"""
InsideInterview AI — Resume Parser
Extracts text from PDF and DOCX files.
"""

import io
from PyPDF2 import PdfReader
from docx import Document


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text content from a PDF file."""
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        return "\n".join(text_parts).strip()
    except Exception as e:
        return f"Error reading PDF: {str(e)}"


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text content from a DOCX file."""
    try:
        doc = Document(io.BytesIO(file_bytes))
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        return "\n".join(text_parts).strip()
    except Exception as e:
        return f"Error reading DOCX: {str(e)}"


def parse_resume(uploaded_file) -> str:
    """
    Parse an uploaded resume file and return extracted text.
    Supports PDF and DOCX formats.
    """
    file_bytes = uploaded_file.read()
    filename = uploaded_file.name.lower()

    if filename.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif filename.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    else:
        return "Unsupported file format. Please upload a PDF or DOCX file."
